import os, json, re, io
from flask import Flask, render_template, request, jsonify, send_file
from PyPDF2 import PdfReader
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

app = Flask(__name__)

API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
APP_PASSWORD = os.environ.get("APP_PASSWORD", "MyOrg2025!Secure")
client = Anthropic(api_key=API_KEY)

# ─── Text Extraction ─────────────────────────────────────
def extract_pdf_text(file_storage):
    reader = PdfReader(file_storage)
    text = ""
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text += t + "\n"
    return text

def extract_html_text(file_storage):
    content = file_storage.read().decode("utf-8", errors="ignore")
    clean = re.sub(r'<[^>]+>', ' ', content)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean

def extract_file_text(file_storage):
    name = file_storage.filename.lower()
    if name.endswith(".pdf"):
        return extract_pdf_text(file_storage)
    elif name.endswith(".html") or name.endswith(".htm"):
        return extract_html_text(file_storage)
    else:
        return file_storage.read().decode("utf-8", errors="ignore")

# ─── Claude AI Call ───────────────────────────────────────
TAX_PROMPT = """IMPORTANT: Return ONLY valid JSON with no additional text, no explanations, no markdown formatting.

Extract the following tax return data for FY {year} as pure JSON:

{{
  "tax_year": "{year}",
  "taxpayer_name": "",
  "summary": {{
    "total_income": 0, "total_deductions": 0, "taxable_income": 0,
    "tax_withheld": 0, "estimated_amount": 0, "is_payable": true
  }},
  "income": {{
    "salary": 0, "interest": 0, "dividends": 0, "rental_gross": 0,
    "rental_net": 0, "capital_gains": 0, "other": 0
  }},
  "capital_gains_details": {{
    "current_year_capital_gain": 0, "current_year_capital_loss": 0,
    "net_capital_gain_loss": 0, "capital_losses_carried_forward": 0,
    "capital_losses_applied": 0, "capital_losses_unapplied": 0,
    "capital_gains_discount_amount": 0, "capital_gains_taxable": 0
  }},
  "deductions": {{
    "work_related_car": 0, "work_related_travel": 0, "work_related_clothing": 0,
    "work_related_self_education": 0, "work_related_other": 0,
    "interest_deductions": 0, "dividend_deductions": 0, "gifts_donations": 0,
    "cost_managing_tax_affairs": 0, "total": 0
  }},
  "rental_properties": [
    {{
      "address": "Full property address", "ownership_percentage": 100, "gross_rent": 0,
      "expenses": {{
        "interest_on_loans": 0, "body_corporate_strata_fees": 0, "council_rates": 0,
        "water_rates": 0, "insurance": 0, "land_tax": 0, "property_agent_fees": 0,
        "repairs_maintenance": 0, "depreciation_div_40": 0, "depreciation_div_43": 0,
        "capital_allowances": 0, "other_rental_expenses": 0, "total_expenses": 0
      }},
      "net_income_loss": 0
    }}
  ],
  "all_deductions_list": {{
    "D1": {{ "description": "Work car expenses", "amount": 0 }},
    "D2": {{ "description": "Work travel expenses", "amount": 0 }},
    "D3": {{ "description": "Work clothing expenses", "amount": 0 }},
    "D4": {{ "description": "Work self-education", "amount": 0 }},
    "D5": {{ "description": "Other work deductions", "amount": 0 }},
    "D6": {{ "description": "Low value pool deduction", "amount": 0 }},
    "D7": {{ "description": "Interest deductions", "amount": 0 }},
    "D8": {{ "description": "Dividend deductions", "amount": 0 }},
    "D9": {{ "description": "Gifts or donations", "amount": 0 }},
    "D10": {{ "description": "Cost of managing tax affairs", "amount": 0 }},
    "D11": {{ "description": "Deductible amount of UPP", "amount": 0 }},
    "D12": {{ "description": "Personal superannuation", "amount": 0 }},
    "D13": {{ "description": "Deduction for project pool", "amount": 0 }},
    "D14": {{ "description": "Forestry managed investment", "amount": 0 }},
    "D15": {{ "description": "Other deductions", "amount": 0 }}
  }}
}}

CRITICAL EXTRACTION RULES:
1. Extract EVERY deduction item with its D-code (D1-D15) and exact amount
2. For rental properties, extract ONLY addresses from the "Rental Properties Schedule"
3. Look for "ESTIMATED AMOUNT PAYABLE" or "REFUND" for estimated_amount
4. Set is_payable to true if payable, false if refund
5. All amounts should be exact as shown in the return
6. Return ONLY valid JSON, no other text"""

def call_claude(pdf_text, year, sop_context=None):
    prompt = TAX_PROMPT.format(year=year)
    prompt += f"\n\nHere is the tax return document text:\n---\n{pdf_text}\n---"
    if sop_context:
        if sop_context.get("prefill"):
            prompt += f"\n\nPREFILL DATA:\n---\n{sop_context['prefill']}\n---"
        if sop_context.get("questionnaire"):
            prompt += f"\n\nCLIENT QUESTIONNAIRE:\n---\n{sop_context['questionnaire']}\n---"
        if sop_context.get("supporting"):
            prompt += f"\n\nSUPPORTING DOCUMENTS:\n---\n{sop_context['supporting']}\n---"
        prompt += "\n\nSOP CROSS-CHECK: Compare return amounts against prefill/questionnaire/supporting docs. Flag discrepancies."
    prompt += "\n\nReturn the extracted data as valid JSON only."

    msg = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )
    txt = msg.content[0].text
    if "```json" in txt:
        txt = txt.split("```json")[1].split("```")[0]
    elif "```" in txt:
        txt = txt.split("```")[1].split("```")[0]
    return json.loads(txt.strip())

# ─── Comparison Logic ─────────────────────────────────────
def validate_rental(d):
    if not isinstance(d.get("rental_properties"), list):
        d["rental_properties"] = []
    d["rental_properties"] = [p for p in d["rental_properties"]
        if p.get("address") and not re.match(r'(?i)full property address|property address', p["address"])
        and (p.get("gross_rent", 0) != 0 or p.get("net_income_loss", 0) != 0)]

def norm_addr(a):
    return re.sub(r'\s+', ' ', a.lower()).strip()

STRIP_WORDS = ["australia","aus","nsw","vic","qld","sa","wa","tas","nt","act",
    "new south wales","victoria","queensland","south australia","western australia",
    "tasmania","northern territory","australian capital territory"]

def core_addr(a):
    s = norm_addr(a)
    s = re.sub(r'\b\d{4}\b', '', s)
    for w in STRIP_WORDS:
        s = re.sub(r'\b' + w + r'\b', '', s, flags=re.I)
    s = re.sub(r',\s*,', ',', s)
    s = re.sub(r',\s*$', '', s)
    s = re.sub(r'^\s*,', '', s)
    return re.sub(r'\s+', ' ', s).strip()

def street_part(a):
    return core_addr(a).split(",")[0].strip()

def word_sim(a, b):
    wa = set(w for w in re.split(r'[\s,]+', a) if len(w) > 1)
    wb = set(w for w in re.split(r'[\s,]+', b) if len(w) > 1)
    if not wa or not wb:
        return 0
    shared = len(wa & wb)
    return shared / max(len(wa), len(wb))

def find_prop(addr, lst):
    n = norm_addr(addr)
    for p in lst:
        if norm_addr(p["address"]) == n:
            return p
    ca = core_addr(addr)
    for p in lst:
        if core_addr(p["address"]) == ca:
            return p
    sp = street_part(addr)
    if len(sp) > 5:
        for p in lst:
            if street_part(p["address"]) == sp:
                return p
    best_score, best = 0, None
    for p in lst:
        score = word_sim(core_addr(addr), core_addr(p["address"]))
        if score > best_score:
            best_score, best = score, p
    if best_score >= 0.7:
        return best
    return None

EXPENSE_KEYS = [
    ("interest_on_loans", "Interest on Loans", True),
    ("body_corporate_strata_fees", "Body Corporate", True),
    ("council_rates", "Council Rates", True),
    ("water_rates", "Water Charges", True),
    ("insurance", "Insurance", True),
    ("land_tax", "Land Tax", False),
    ("property_agent_fees", "Agent Fees", False),
    ("repairs_maintenance", "Repairs & Maintenance", False),
    ("depreciation_div_40", "Depreciation Div-40", False),
    ("depreciation_div_43", "Capital Works (Div 43)", True),
    ("capital_allowances", "Capital Allowances", False),
    ("other_rental_expenses", "Sundry Expenses", False),
]

def compare_returns(curr, prev):
    alerts = []
    tax_change = curr["summary"]["estimated_amount"] - prev["summary"]["estimated_amount"]
    income_change = curr["summary"]["total_income"] - prev["summary"]["total_income"]
    ded_change = curr["summary"]["total_deductions"] - prev["summary"]["total_deductions"]

    # Tax position alerts
    if not prev["summary"]["is_payable"] and not curr["summary"]["is_payable"]:
        rc = abs(tax_change)
        if rc > 0:
            alerts.append({"type": "success" if tax_change > 0 else "warning",
                "message": f"Refund {'increased' if tax_change > 0 else 'decreased'} by ${rc:,.0f}"})
    elif prev["summary"]["is_payable"] != curr["summary"]["is_payable"]:
        good = prev["summary"]["is_payable"] and not curr["summary"]["is_payable"]
        alerts.append({"type": "success" if good else "critical",
            "message": f"Tax position changed from {'payable' if prev['summary']['is_payable'] else 'refund'} to {'payable' if curr['summary']['is_payable'] else 'refund'}"})
    elif abs(tax_change) > 100:
        alerts.append({"type": "critical" if abs(tax_change) > 5000 else "warning",
            "message": f"Tax {'payable' if curr['summary']['is_payable'] else 'refund'} changed by ${abs(tax_change):,.0f}"})

    # Salary
    sal_change = (curr["income"].get("salary",0) or 0) - (prev["income"].get("salary",0) or 0)
    if abs(sal_change) >= 1000:
        alerts.append({"type": "warning" if abs(sal_change)>10000 else "info",
            "message": f"Salary {'increased' if sal_change>0 else 'decreased'} by ${abs(sal_change):,.0f}"})

    # Deductions comparison
    ded_breakdown = []
    missing_deds = []
    new_deds = []
    p_ded = prev.get("all_deductions_list", {})
    c_ded = curr.get("all_deductions_list", {})
    for dc in p_ded:
        pa = p_ded[dc].get("amount", 0) or 0
        ca = (c_ded.get(dc, {}).get("amount", 0)) or 0
        desc = c_ded.get(dc, p_ded[dc]).get("description", dc)
        if pa > 0 or ca > 0:
            ded_breakdown.append({"code": dc, "description": desc, "fy_prev": pa, "fy_curr": ca, "change": ca - pa})
        if pa > 0 and ca == 0:
            missing_deds.append({"code": dc, "description": desc, "previous_amount": pa})
            alerts.append({"type": "critical", "message": f"MISSING DEDUCTION: {desc} (${pa:,.0f} last year, $0 this year)"})
        elif pa == 0 and ca > 0:
            new_deds.append({"code": dc, "description": desc, "current_amount": ca})
            alerts.append({"type": "info", "message": f"NEW DEDUCTION: {desc} (${ca:,.0f})"})

    # Rental comparison
    rental_details = []
    prev_props = prev.get("rental_properties", [])
    curr_props = curr.get("rental_properties", [])
    for pp in prev_props:
        cp = find_prop(pp["address"], curr_props)
        if cp:
            expenses = []
            for key, name, critical in EXPENSE_KEYS:
                pe = (pp.get("expenses", {}).get(key, 0)) or 0
                ce = (cp.get("expenses", {}).get(key, 0)) or 0
                status = "missing" if pe > 0 and ce == 0 else "new" if pe == 0 and ce > 0 else "unchanged" if ce-pe == 0 else "increased" if ce > pe else "decreased"
                expenses.append({"type": name, "fy_prev": pe, "fy_curr": ce, "change": ce-pe, "status": status})
                if pe > 0 and ce == 0:
                    alerts.append({"type": "critical" if critical else "warning",
                        "message": f'Property "{pp["address"]}": {name} missing (was ${pe:,.0f})'})
            rental_details.append({
                "address": pp["address"], "status": "continuing",
                "fy_prev_gross": pp.get("gross_rent", 0) or 0,
                "fy_curr_gross": cp.get("gross_rent", 0) or 0,
                "fy_prev_net": pp.get("net_income_loss", 0),
                "fy_curr_net": cp.get("net_income_loss", 0),
                "change": (cp.get("net_income_loss",0)) - (pp.get("net_income_loss",0)),
                "all_expenses": expenses
            })
        else:
            tot = sum((pp.get("expenses",{}).get(k,0) or 0) for k,_,_ in EXPENSE_KEYS)
            if abs(pp.get("net_income_loss",0)) > 100 or tot > 500:
                alerts.append({"type": "warning", "message": f'Property REMOVED: "{pp["address"]}"'})
                rental_details.append({"address": pp["address"], "status": "removed",
                    "fy_prev_gross": pp.get("gross_rent",0), "fy_curr_gross": 0,
                    "fy_prev_net": pp.get("net_income_loss",0), "fy_curr_net": 0,
                    "change": -(pp.get("net_income_loss",0) or 0), "all_expenses": []})

    for cp in curr_props:
        if not find_prop(cp["address"], prev_props):
            alerts.append({"type": "info", "message": f'NEW Property: "{cp["address"]}"'})
            rental_details.append({"address": cp["address"], "status": "new",
                "fy_prev_gross": 0, "fy_curr_gross": cp.get("gross_rent",0),
                "fy_prev_net": 0, "fy_curr_net": cp.get("net_income_loss",0),
                "change": cp.get("net_income_loss",0) or 0, "all_expenses": []})

    if not alerts:
        alerts.append({"type": "success", "message": "No significant changes detected"})
    order = {"critical": 1, "warning": 2, "info": 3, "success": 4}
    alerts.sort(key=lambda a: order.get(a["type"], 5))

    # Income breakdown
    income_breakdown = {}
    for k, label in [("salary","Salary"),("interest","Interest"),("dividends","Dividends"),
                      ("capital_gains","Capital Gains"),("other","Other")]:
        p = prev["income"].get(k, 0) or 0
        c = curr["income"].get(k, 0) or 0
        if p or c:
            income_breakdown[k] = {"label": label, "fy_prev": p, "fy_curr": c, "change": c-p}
    rp = prev["income"].get("rental_net", 0) or 0
    rc = curr["income"].get("rental_net", 0) or 0
    if rp or rc:
        income_breakdown["rental"] = {"label": "Rental", "fy_prev": rp, "fy_curr": rc, "change": rc-rp}
    ti_p = prev["summary"]["total_income"]
    ti_c = curr["summary"]["total_income"]
    income_breakdown["total"] = {"label": "TOTAL", "fy_prev": ti_p, "fy_curr": ti_c, "change": ti_c-ti_p}

    return {
        "previous": {"year": prev["tax_year"], "amount": prev["summary"]["estimated_amount"],
            "type": "Payable" if prev["summary"]["is_payable"] else "Refund",
            "totalIncome": prev["summary"]["total_income"], "deductions": prev["summary"]["total_deductions"],
            "tax_withheld": prev["summary"].get("tax_withheld", 0)},
        "current": {"year": curr["tax_year"], "amount": curr["summary"]["estimated_amount"],
            "type": "Payable" if curr["summary"]["is_payable"] else "Refund",
            "totalIncome": curr["summary"]["total_income"], "deductions": curr["summary"]["total_deductions"],
            "tax_withheld": curr["summary"].get("tax_withheld", 0)},
        "changes": {"tax": tax_change, "income": income_change,
            "salary": sal_change, "rental": rc - rp, "deductions": ded_change},
        "deduction_analysis": {"detailed_breakdown": ded_breakdown, "missing": missing_deds,
            "new_this": new_deds, "total_change": ded_change},
        "rental_breakdown": {"property_details": rental_details},
        "income_breakdown": income_breakdown,
        "alerts": alerts,
        "_rawPrev": prev, "_rawCurr": curr
    }

# ─── DOCX Report Generator ───────────────────────────────
def generate_docx(result):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    prev_raw = result["_rawPrev"]
    curr_raw = result["_rawCurr"]
    name = curr_raw.get("taxpayer_name") or prev_raw.get("taxpayer_name") or "Client"
    fy_prev = result["previous"]["year"]
    fy_curr = result["current"]["year"]

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TAX RETURN ANALYSIS REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{name} \u2014 FY{fy_prev} vs FY{fy_curr}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    from datetime import date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {date.today().strftime('%d/%m/%Y')} | ACEBIZ \u2014 C&S Partners Australia Pty Ltd")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Tax Position Summary ──
    h = doc.add_heading("Tax Position Summary", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    summary_rows = [
        ("Salary/Wages", prev_raw["income"].get("salary",0), curr_raw["income"].get("salary",0)),
        ("Interest Income", prev_raw["income"].get("interest",0), curr_raw["income"].get("interest",0)),
        ("Dividend Income", prev_raw["income"].get("dividends",0), curr_raw["income"].get("dividends",0)),
        ("Net Rental Income", prev_raw["income"].get("rental_net",0), curr_raw["income"].get("rental_net",0)),
        ("Capital Gains", prev_raw["income"].get("capital_gains",0), curr_raw["income"].get("capital_gains",0)),
        ("Other Income", prev_raw["income"].get("other",0), curr_raw["income"].get("other",0)),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["", f"FY{fy_prev}", f"FY{fy_curr}", "Change"]):
        hdr[i].text = txt
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for label, pv, cv in summary_rows:
        pv = pv or 0; cv = cv or 0
        if pv == 0 and cv == 0:
            continue
        ch = cv - pv
        flag = " !!" if abs(ch) > 1000 else ""
        row = tbl.add_row().cells
        row[0].text = label
        row[1].text = f"${pv:,.0f}"
        row[2].text = f"${cv:,.0f}"
        row[3].text = f"{'+'if ch>0 else ''}${ch:,.0f}{flag}"
        for c in row[1:]:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Deductions section
    da = result["deduction_analysis"]
    if da["detailed_breakdown"]:
        row = tbl.add_row().cells
        row[0].text = "Total Deductions"
        row[1].text = f"${result['previous']['deductions']:,.0f}"
        row[2].text = f"${result['current']['deductions']:,.0f}"
        row[3].text = f"{'+'if result['changes']['deductions']>0 else ''}${result['changes']['deductions']:,.0f}"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                if c != row[0]:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for d in da["detailed_breakdown"]:
            if d["fy_prev"] == 0 and d["fy_curr"] == 0:
                continue
            flag = " !!" if d["fy_prev"] > 0 and d["fy_curr"] == 0 else ""
            row = tbl.add_row().cells
            row[0].text = f"  {d['code']} {d['description']}"
            row[1].text = f"${d['fy_prev']:,.0f}"
            row[2].text = f"${d['fy_curr']:,.0f}"
            row[3].text = f"{'+'if d['change']>0 else ''}${d['change']:,.0f}{flag}"
            for c in row[1:]:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Tax withheld
    tw_p = result["previous"].get("tax_withheld", 0) or 0
    tw_c = result["current"].get("tax_withheld", 0) or 0
    row = tbl.add_row().cells
    row[0].text = "Tax Withheld"
    row[1].text = f"${tw_p:,.0f}"
    row[2].text = f"${tw_c:,.0f}"
    row[3].text = f"{'+'if tw_c-tw_p>0 else ''}${tw_c-tw_p:,.0f}"
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
            if c != row[0]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Est tax
    row = tbl.add_row().cells
    row[0].text = f"Est. Tax {result['previous']['type']}"
    row[1].text = f"${abs(result['previous']['amount']):,.0f}"
    row[2].text = f"${abs(result['current']['amount']):,.0f}"
    ch = result["changes"]["tax"]
    row[3].text = f"{'+'if ch>0 else ''}${ch:,.0f}"
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
            if c != row[0]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Critical Findings ──
    h = doc.add_heading("Critical Findings & SOP Checks", level=2)
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    icons = {"critical": "\U0001F6A8", "warning": "\u26A0\uFE0F", "info": "\u2139\uFE0F", "success": "\u2705"}
    ftbl = doc.add_table(rows=0, cols=2)
    ftbl.style = 'Table Grid'
    for a in result["alerts"]:
        row = ftbl.add_row().cells
        row[0].text = icons.get(a["type"], "")
        row[1].text = a["message"]
        row[0].width = Cm(1.2)
        for p in row[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Rental Property Comparison ──
    rental = result["rental_breakdown"]["property_details"]
    if rental:
        h = doc.add_heading(f"Rental Property Comparison ({len(rental)} Properties)", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        for prop in rental:
            doc.add_heading(prop["address"], level=3)
            ptbl = doc.add_table(rows=1, cols=5)
            ptbl.style = 'Table Grid'
            for i, txt in enumerate(["Item", f"FY{fy_prev}", f"FY{fy_curr}", "Change", "Flag"]):
                ptbl.rows[0].cells[i].text = txt
                for p in ptbl.rows[0].cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

            # Gross rent
            gp = prop.get("fy_prev_gross", 0) or 0
            gc = prop.get("fy_curr_gross", 0) or 0
            row = ptbl.add_row().cells
            row[0].text = "Gross Rent (Income)"
            row[1].text = f"${gp:,.0f}"
            row[2].text = f"${gc:,.0f}"
            row[3].text = f"{'+'if gc-gp>0 else ''}${gc-gp:,.0f}"
            row[4].text = ""
            for c in row[1:]:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Expenses
            total_prev_exp = 0
            total_curr_exp = 0
            for e in prop.get("all_expenses", []):
                flag = "MISSING" if e["status"] == "missing" else ("CHECK" if abs(e["change"]) > 1000 else "")
                row = ptbl.add_row().cells
                row[0].text = e["type"]
                row[1].text = f"${e['fy_prev']:,.0f}"
                row[2].text = f"${e['fy_curr']:,.0f}"
                row[3].text = f"{'+'if e['change']>0 else ''}${e['change']:,.0f}"
                row[4].text = flag
                total_prev_exp += e["fy_prev"]
                total_curr_exp += e["fy_curr"]
                for c in row[1:]:
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Total expenses
            row = ptbl.add_row().cells
            row[0].text = "Total Expenses"
            row[1].text = f"${total_prev_exp:,.0f}"
            row[2].text = f"${total_curr_exp:,.0f}"
            ch = total_curr_exp - total_prev_exp
            row[3].text = f"{'+'if ch>0 else ''}${ch:,.0f}"
            row[4].text = ""
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
                    if c != row[0]:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Net result
            row = ptbl.add_row().cells
            row[0].text = "Net Result"
            row[1].text = f"${prop['fy_prev_net']:,.0f}"
            row[2].text = f"${prop['fy_curr_net']:,.0f}"
            row[3].text = f"{'+'if prop['change']>0 else ''}${prop['change']:,.0f}"
            row[4].text = ""
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.bold = True
                    if c != row[0]:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("ACEBIZ \u2014 C&S Partners Australia Pty Ltd | Tax Agent No. 24742337")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p = doc.add_paragraph()
    run = p.add_run("204/11 Solent Circuit, Norwest NSW 2153")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc

# ─── Routes ───────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    if data.get("password") == APP_PASSWORD:
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Incorrect password"})

@app.route("/analyze", methods=["POST"])
def analyze():
    prev_file = request.files.get("prevFile")
    curr_file = request.files.get("currFile")
    prev_year = request.form.get("prevYear", "2024")
    curr_year = request.form.get("currYear", "2025")

    if not prev_file or not curr_file:
        return jsonify({"error": "Both tax return PDFs are required"}), 400

    sop_context = {}
    prefill = request.files.get("prefillFile")
    quest = request.files.get("questFile")
    supp_files = request.files.getlist("suppFile")

    prev_text = extract_pdf_text(prev_file)
    curr_text = extract_pdf_text(curr_file)

    if prefill:
        sop_context["prefill"] = extract_file_text(prefill)
    if quest:
        sop_context["questionnaire"] = extract_file_text(quest)
    if supp_files and supp_files[0].filename:
        supp_text = ""
        for sf in supp_files:
            if sf.filename.lower().endswith(".pdf"):
                supp_text += f"\n[{sf.filename}]:\n{extract_pdf_text(sf)}\n"
            else:
                supp_text += f"\n[{sf.filename}]: (image file)\n"
        sop_context["supporting"] = supp_text

    sop = sop_context if sop_context else None

    prev_data = call_claude(prev_text, prev_year, sop)
    curr_data = call_claude(curr_text, curr_year, sop)
    validate_rental(prev_data)
    validate_rental(curr_data)

    result = compare_returns(curr_data, prev_data)
    # Return everything including raw data (needed for report generation)
    return jsonify(result)

@app.route("/download-report", methods=["POST"])
def download_report():
    result = request.get_json()
    if not result or "_rawPrev" not in result:
        return "No analysis data provided", 400
    doc = generate_docx(result)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    name = result["_rawCurr"].get("taxpayer_name") or result["_rawPrev"].get("taxpayer_name") or "Client"
    fname = f"{name.replace(' ', '_')}_Tax_Analysis_FY{result['previous']['year']}_vs_FY{result['current']['year']}.docx"
    return send_file(buf, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True, port=5000)
